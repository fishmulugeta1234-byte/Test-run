import os
import re
import shutil
import subprocess
import threading
from datetime import datetime
from http.server import HTTPServer, BaseHTTPRequestHandler
from pathlib import Path

from openpyxl import load_workbook
from docx import Document
from telegram import Update
from telegram.ext import (
    Application,
    CommandHandler,
    MessageHandler,
    ContextTypes,
    filters,
)

# ============================================================
# SIMON ORIGIN TRANSFORMATION
# PRIVATE ADMIN PLAN GENERATOR
# ============================================================
# You collect the client's information yourself.
# The bot then:
#   1. Calculates calories/macros from your Excel reference.
#   2. Builds a workout from the selected equipment/experience.
#   3. Fills your existing Workout DOCX template.
#   4. Fills your existing Nutrition DOCX template.
#   5. Converts both DOCX files to PDF using LibreOffice.
#   6. Sends both PDFs back to YOU.
#
# Required files in the same GitHub repository:
#   Bot.py
#   Calorie_Macro_Calculator.xlsx
#   AbyssinicaSIL-Regular.ttf (kept for your existing project;
#       the DOCX templates already contain their own formatting)
#   30-Day_Workout_Blueprint_TEMPLATE.docx
#   30-Day_Nutrition_Blueprint_TEMPLATE.docx
#
# Required Render environment variables:
#   BOT_TOKEN
#   ADMIN_ID
#
# ============================================================

BASE_DIR = Path(__file__).resolve().parent

CALCULATOR_FILE = BASE_DIR / "Calorie_Macro_Calculator.xlsx"
WORKOUT_TEMPLATE = BASE_DIR / "30-Day_Workout_Blueprint_TEMPLATE.docx"
NUTRITION_TEMPLATE = BASE_DIR / "30-Day_Nutrition_Blueprint_TEMPLATE.docx"

OUTPUT_DIR = BASE_DIR / os.getenv("OUTPUT_DIR", "generated_plans")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

BOT_TOKEN = os.getenv("BOT_TOKEN", "").strip()
ADMIN_ID_RAW = os.getenv("ADMIN_ID", "").strip()

if not BOT_TOKEN:
    raise RuntimeError("BOT_TOKEN is missing from Render Environment Variables.")

if not ADMIN_ID_RAW.isdigit():
    raise RuntimeError("ADMIN_ID must be your numeric Telegram user ID.")

ADMIN_ID = int(ADMIN_ID_RAW)


# ============================================================
# 1. RENDER HEALTH CHECK
# ============================================================

class HealthCheck(BaseHTTPRequestHandler):
    def do_GET(self):
        self.send_response(200)
        self.send_header("Content-Type", "text/plain; charset=utf-8")
        self.end_headers()
        self.wfile.write(b"Simon Origin Plan Generator is active.")

    def log_message(self, format, *args):
        return


def start_health_check():
    port = int(os.getenv("PORT", "8080"))
    server = HTTPServer(("0.0.0.0", port), HealthCheck)
    server.serve_forever()


threading.Thread(target=start_health_check, daemon=True).start()


# ============================================================
# 2. FILE CHECKS
# ============================================================

def check_required_files():
    missing = []

    for file_path in (
        CALCULATOR_FILE,
        WORKOUT_TEMPLATE,
        NUTRITION_TEMPLATE,
    ):
        if not file_path.exists():
            missing.append(file_path.name)

    if missing:
        raise RuntimeError(
            "Missing required files: " + ", ".join(missing)
        )


check_required_files()


# ============================================================
# 3. ADMIN SECURITY
# ============================================================

def is_admin(update: Update) -> bool:
    user = update.effective_user
    return bool(user and user.id == ADMIN_ID)


async def deny(update: Update):
    if update.effective_message:
        await update.effective_message.reply_text(
            "⛔ This is a private Simon Origin coaching tool."
        )


# ============================================================
# 4. READ ACTIVITY + GOAL VALUES FROM YOUR EXCEL FILE
# ============================================================

def load_calculator_reference():
    workbook = load_workbook(CALCULATOR_FILE, data_only=True)

    if "Calculator" not in workbook.sheetnames:
        raise RuntimeError(
            "The Excel file must contain a sheet named 'Calculator'."
        )

    sheet = workbook["Calculator"]

    activity = {}
    for row in range(9, 13):
        name = sheet[f"E{row}"].value
        multiplier = sheet[f"F{row}"].value

        if name is not None and multiplier is not None:
            activity[str(name).strip()] = float(multiplier)

    goals = {}
    for row in range(9, 12):
        name = sheet[f"G{row}"].value
        adjustment = sheet[f"H{row}"].value

        if name is not None and adjustment is not None:
            goals[str(name).strip()] = float(adjustment)

    if not activity:
        raise RuntimeError("No activity multipliers found in Excel.")

    if not goals:
        raise RuntimeError("No goal adjustments found in Excel.")

    return activity, goals


ACTIVITY_MULTIPLIERS, GOAL_ADJUSTMENTS = load_calculator_reference()


# ============================================================
# 5. NUTRITION CALCULATION
# ============================================================

def calculate_macros(data):
    gender = data["gender"].strip().lower()
    age = float(data["age"])
    height = float(data["height"])
    weight = float(data["weight"])

    activity = data["activity"]
    goal = data["goal"]

    # Mifflin-St Jeor
    if gender == "male":
        bmr = (10 * weight) + (6.25 * height) - (5 * age) + 5
    else:
        bmr = (10 * weight) + (6.25 * height) - (5 * age) - 161

    if activity not in ACTIVITY_MULTIPLIERS:
        raise ValueError(f"Unknown activity level: {activity}")

    if goal not in GOAL_ADJUSTMENTS:
        raise ValueError(f"Unknown goal: {goal}")

    tdee = bmr * ACTIVITY_MULTIPLIERS[activity]

    # Uses the goal adjustment from your Excel reference.
    calories = tdee * (1 + GOAL_ADJUSTMENTS[goal])

    # Your calculator structure:
    # Protein = 2.0 g/kg
    # Fat = 25% calories
    # Carbs = remaining calories
    protein = weight * 2.0
    protein_calories = protein * 4

    fat_calories = calories * 0.25
    fats = fat_calories / 9

    carb_calories = calories - protein_calories - fat_calories
    carbs = max(0, carb_calories / 4)

    # Your blueprint uses 0.05 L/kg.
    water = weight * 0.05

    return {
        "bmr": round(bmr),
        "tdee": round(tdee),
        "calories": round(calories),
        "weekly_calories": round(calories * 7),
        "protein": round(protein),
        "carbs": round(carbs),
        "fats": round(fats),
        "water": round(water, 1),
    }


# ============================================================
# 6. WORKOUT LIBRARY
# ============================================================
# The supplied workout template contains [Exercise] placeholders,
# so these are the exercise selections used to populate them.
# Review every generated plan before sending it to a client.

GYM_PLAN = [
    [
        ("Machine Chest Press / የደረት ማሽን", "3", "8-12", "90s", "Controlled reps / በቁጥጥር"),
        ("Lat Pulldown / ላት ፑልዳውን", "3", "8-12", "90s", "Full range / ሙሉ ክልል"),
        ("Seated Cable Row / ሲቲድ ኬብል ሮ", "3", "10-12", "90s", "Neutral spine / ጀርባ ቀጥ"),
        ("Dumbbell Shoulder Press / ዳምቤል የትከሻ ፕሬስ", "2", "10-12", "75s", "Light-moderate load"),
        ("Cable Triceps Pressdown / ትራይሴፕስ", "2", "12-15", "60s", "Smooth movement"),
    ],
    [
        ("Goblet Squat / ጎብሌት ስኳት", "3", "8-12", "90s", "Controlled depth"),
        ("Romanian Deadlift / ሮማኒያን ዴድሊፍት", "3", "8-12", "90s", "Hinge at hips"),
        ("Leg Press / ሌግ ፕሬስ", "3", "10-12", "90s", "Do not lock knees"),
        ("Standing Calf Raise / የእግር ጫፍ", "3", "12-15", "60s", "Full range"),
        ("Dead Bug / ዴድ ባግ", "3", "8-12/side", "45s", "Brace core"),
    ],
    [
        ("Brisk Walk / ፈጣን መራመድ", "1", "20-30 min", "—", "Easy-moderate pace"),
        ("Cat-Cow / ካት-ካው", "2", "8-10", "30s", "Gentle mobility"),
        ("Hip Flexor Stretch / የዳሌ መዘርጋት", "2", "30s/side", "20s", "Gentle stretch"),
        ("Thoracic Rotation / የላይኛው ጀርባ", "2", "8/side", "30s", "Slow movement"),
        ("Child's Pose / የልጅ አቀማመጥ", "2", "30-45s", "20s", "Relaxed breathing"),
    ],
    [
        ("Incline Dumbbell Press / ኢንክላይን ዳምቤል", "3", "8-12", "90s", "Controlled"),
        ("Dumbbell Lateral Raise / የትከሻ ጎን", "3", "12-15", "60s", "Light weight"),
        ("Cable Chest Fly / የደረት ፍላይ", "2", "10-15", "60s", "Do not over-stretch"),
        ("Rope Triceps Pressdown / ትራይሴፕስ", "3", "10-15", "60s", "Full extension"),
        ("Plank / ፕላንክ", "3", "20-45s", "45s", "Neutral spine"),
    ],
    [
        ("Assisted Pull-Up / የተረዳ ፑል አፕ", "3", "6-10", "90s", "Controlled"),
        ("Seated Cable Row / ሲቲድ ሮ", "3", "8-12", "90s", "Squeeze back"),
        ("Dumbbell Curl / ቢሴፕስ ከርል", "3", "10-15", "60s", "No swinging"),
        ("Face Pull / ፌስ ፑል", "3", "12-15", "60s", "Shoulder control"),
        ("Back Extension / የጀርባ ኤክስቴንሽን", "2", "10-15", "60s", "Neutral spine"),
    ],
    [
        ("Leg Press / ሌግ ፕሬስ", "3", "10-12", "90s", "Controlled"),
        ("Dumbbell Romanian Deadlift / ዳምቤል RDL", "3", "8-12", "90s", "Hips back"),
        ("Walking Lunge / የመራመጃ ላንጅ", "2", "10/leg", "75s", "Stable steps"),
        ("Leg Curl / ሌግ ከርል", "3", "10-15", "60s", "Controlled"),
        ("Calf Raise / ካልፍ ሬዝ", "3", "12-15", "60s", "Full range"),
    ],
    [
        ("Rest / እረፍት", "—", "—", "—", "Optional easy walk"),
        ("Easy Walk / ቀላል መራመድ", "1", "15-30 min", "—", "Optional"),
        ("Gentle Stretch / ቀላል መዘርጋት", "1", "5-10 min", "—", "Optional"),
        ("Breathing / የመተንፈስ ልምምድ", "1", "3-5 min", "—", "Relax"),
        ("Recovery / ማገገሚያ", "—", "—", "—", "Prioritize sleep"),
    ],
]

HOME_PLAN = [
    [
        ("Incline Push-Up / ኢንክላይን ፑሽ አፕ", "3", "8-15", "60s", "Use wall/table if needed"),
        ("Backpack Row / ቦርሳ ሮ", "3", "10-15", "60s", "Controlled"),
        ("Bodyweight Squat / የሰውነት ስኳት", "3", "10-15", "60s", "Knees track toes"),
        ("Glute Bridge / ግሉት ብሪጅ", "3", "12-15", "60s", "Squeeze glutes"),
        ("Plank / ፕላንክ", "3", "20-40s", "45s", "Brace core"),
    ],
    [
        ("Bodyweight Squat / ስኳት", "3", "10-15", "60s", "Controlled"),
        ("Reverse Lunge / ሪቨርስ ላንጅ", "3", "8-12/leg", "60s", "Stable"),
        ("Glute Bridge / ግሉት ብሪጅ", "3", "12-15", "60s", "Pause at top"),
        ("Calf Raise / ካልፍ ሬዝ", "3", "15-20", "45s", "Full range"),
        ("Dead Bug / ዴድ ባግ", "3", "8-12/side", "45s", "Brace core"),
    ],
    [
        ("Brisk Walk / ፈጣን መራመድ", "1", "20-30 min", "—", "Easy-moderate pace"),
        ("Cat-Cow / ካት-ካው", "2", "8-10", "30s", "Gentle"),
        ("Hip Flexor Stretch / የዳሌ መዘርጋት", "2", "30s/side", "20s", "Gentle"),
        ("Thoracic Rotation / የላይኛው ጀርባ", "2", "8/side", "30s", "Slow"),
        ("Child's Pose / የልጅ አቀማመጥ", "2", "30-45s", "20s", "Relax"),
    ],
    [
        ("Push-Up / ፑሽ አፕ", "3", "8-15", "60s", "Modify as needed"),
        ("Pike Push-Up / ፓይክ ፑሽ አፕ", "3", "6-12", "60s", "Controlled"),
        ("Backpack Row / ቦርሳ ሮ", "3", "10-15", "60s", "Pull to ribs"),
        ("Chair Triceps Dip / ትራይሴፕስ", "2", "8-12", "60s", "Stable chair"),
        ("Plank / ፕላንክ", "3", "20-45s", "45s", "Neutral spine"),
    ],
    [
        ("Backpack Row / ቦርሳ ሮ", "3", "10-15", "60s", "Controlled"),
        ("Reverse Snow Angel / ሪቨርስ ስኖ አንጀል", "3", "10-15", "45s", "Slow"),
        ("Backpack Curl / ቦርሳ ከርል", "2", "10-15", "60s", "No swinging"),
        ("Bird Dog / በርድ ዶግ", "3", "8-12/side", "45s", "Stable hips"),
        ("Superman / ሱፐርማን", "2", "8-12", "45s", "Gentle range"),
    ],
    [
        ("Bodyweight Squat / ስኳት", "3", "10-15", "60s", "Controlled"),
        ("Reverse Lunge / ሪቨርስ ላንጅ", "3", "8-12/leg", "60s", "Stable"),
        ("Single-Leg Glute Bridge / አንድ እግር ግሉት", "2", "8-12/leg", "60s", "Control"),
        ("Calf Raise / ካልፍ ሬዝ", "3", "15-20", "45s", "Full range"),
        ("Wall Sit / ዎል ሲት", "2", "30-45s", "45s", "Comfortable depth"),
    ],
    [
        ("Rest / እረፍት", "—", "—", "—", "Optional easy walk"),
        ("Easy Walk / ቀላል መራመድ", "1", "15-30 min", "—", "Optional"),
        ("Gentle Stretch / ቀላል መዘርጋት", "1", "5-10 min", "—", "Optional"),
        ("Breathing / መተንፈስ", "1", "3-5 min", "—", "Relax"),
        ("Recovery / ማገገሚያ", "—", "—", "—", "Prioritize sleep"),
    ],
]


def adjust_volume(plan, experience):
    if experience == "Beginner":
        return plan

    add_sets = 1 if experience == "Intermediate" else 2
    adjusted = []

    for day in plan:
        new_day = []

        for exercise in day:
            name, sets, reps, rest, notes = exercise

            if sets.isdigit() and "Rest /" not in name:
                sets = str(int(sets) + add_sets)

            new_day.append((name, sets, reps, rest, notes))

        adjusted.append(new_day)

    return adjusted


# ============================================================
# 7. MEAL LIBRARY
# ============================================================
# The supplied nutrition template contains [Add food] placeholders,
# so it does not provide exact foods/portion sizes. These are starter
# Ethiopian/local-food combinations. Review portions and substitutions
# before sending the plan.

MEAL_LIBRARY = [
    "Eggs + oats + banana / እንቁላል + ኦትስ + ሙዝ",
    "Injera + shiro + salad / እንጀራ + ሽሮ + ሰላጣ",
    "Chicken + rice + vegetables / ዶሮ + ሩዝ + አትክልት",
    "Lentils + injera + vegetables / ምስር + እንጀራ + አትክልት",
    "Eggs + injera + vegetables / እንቁላል + እንጀራ + አትክልት",
    "Fish + potatoes + salad / ዓሳ + ድንች + ሰላጣ",
    "Beef + rice + vegetables / ስጋ + ሩዝ + አትክልት",
    "Ergo/yogurt + fruit + oats / እርጎ + ፍራፍሬ + ኦትስ",
]

MEAL_PATTERNS = [
    [0, 2, 1, 7],
    [4, 5, 3, 7],
    [0, 6, 3, 7],
    [1, 2, 5, 7],
    [4, 2, 3, 7],
    [0, 5, 1, 7],
    [4, 6, 3, 7],
]


# ============================================================
# 8. DOCX HELPERS
# ============================================================

def set_cell_text(cell, text):
    paragraph = cell.paragraphs[0]

    for run in paragraph.runs:
        run.text = ""

    if paragraph.runs:
        paragraph.runs[0].text = str(text)
    else:
        paragraph.add_run(str(text))


def replace_in_paragraph(paragraph, replacements):
    original = paragraph.text
    updated = original

    for old, new in replacements.items():
        updated = updated.replace(old, str(new))

    if updated == original:
        return

    for run in paragraph.runs:
        run.text = ""

    if paragraph.runs:
        paragraph.runs[0].text = updated
    else:
        paragraph.add_run(updated)


def replace_everywhere(doc, replacements):
    def process_table(table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements)

                for nested_table in cell.tables:
                    process_table(nested_table)

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        process_table(table)

    for section in doc.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                replace_in_paragraph(paragraph, replacements)

            for table in container.tables:
                process_table(table)


def fill_client_snapshot(table, data):
    for row in table.rows:
        if len(row.cells) < 2:
            continue

        label = row.cells[0].text.lower()
        value = None

        if "full name" in label or "ሙሉ ስም" in label:
            value = data["name"]

        elif "gender" in label or "ጾታ" in label:
            value = data["gender"]

        elif "age" in label or "እድሜ" in label:
            value = data["age"]

        elif "height" in label or "ቁመት" in label:
            value = f"{data['height']} cm"

        elif "weight" in label or "ክብደት" in label:
            value = f"{data['weight']} kg"

        elif "primary goal" in label or "ዋና ግብ" in label:
            value = data["goal"]

        elif "activity level" in label or "የእንቅስቃሴ" in label:
            value = data["activity"]

        elif "fitness level" in label or "ብቃት" in label:
            value = data["experience"]

        elif "training experience" in label or "ልምድ" in label:
            value = data["experience"]

        elif "equipment" in label or "መሳሪያ" in label:
            value = data["equipment"]

        elif "main obstacle" in label or "ተግዳሮት" in label:
            value = data.get("obstacle", "Consistency")

        elif (
            "restriction" in label
            or "allerg" in label
            or "ገደቦች" in label
            or "አለርጂ" in label
        ):
            value = data.get("restrictions", "None")

        elif (
            "health" in label
            or "injur" in label
            or "ጤና" in label
            or "ጉዳት" in label
        ):
            value = data.get("injuries", "None")

        elif "eating pattern" in label or "አመጋገብ ስርዓት" in label:
            value = data.get("meals", "4 meals/day")

        elif "training preference" in label or "ምርጫ" in label:
            value = data.get("training_preference", "General Fitness")

        if value is not None:
            set_cell_text(row.cells[1], value)


# ============================================================
# 9. FILL NUTRITION DOCX
# ============================================================

def fill_nutrition_doc(data, macros, output_docx):
    doc = Document(NUTRITION_TEMPLATE)

    replacements = {
        "[INSERT CLIENT NAME]": data["name"],
        "[INSERT GENDER]": data["gender"],
        "[INSERT AGE]": data["age"],
        "[INSERT HEIGHT]": data["height"],
        "[INSERT WEIGHT]": data["weight"],
        "[Fat Loss / Muscle Gain / Maintenance]": data["goal"],
        "[Sedentary / Lightly Active / Moderate / Very Active]": data["activity"],
        "[Beginner / Intermediate / Advanced]": data["experience"],
        "[Consistency]": data.get("obstacle", "Consistency"),
        "[INSERT RESTRICTIONS or 'None']": data.get("restrictions", "None"),
        "[INSERT INJURIES or 'None']": data.get("injuries", "None"),
        "[4 meals/day]": data.get("meals", "4 meals/day"),
        "[CALCULATED TARGET]": f"{macros['calories']:,}",
        "[DAILY TARGET x 7]": f"{macros['weekly_calories']:,}",
    }

    replace_everywhere(doc, replacements)

    # Table 1 = client snapshot.
    fill_client_snapshot(doc.tables[1], data)

    # Table 2 = calorie/macro targets.
    macro_table = doc.tables[2]

    set_cell_text(
        macro_table.rows[0].cells[1],
        f"{macros['calories']:,} kcal / day"
    )

    set_cell_text(
        macro_table.rows[1].cells[1],
        "Mifflin-St Jeor BMR × Activity Multiplier, adjusted for goal"
    )

    set_cell_text(
        macro_table.rows[2].cells[1],
        f"{macros['protein']} g  (~2.0g/kg body weight)"
    )

    set_cell_text(
        macro_table.rows[3].cells[1],
        f"{macros['carbs']} g  (periodized by training day)"
    )

    set_cell_text(
        macro_table.rows[4].cells[1],
        f"{macros['fats']} g"
    )

    set_cell_text(
        macro_table.rows[5].cells[1],
        f"{macros['water']} Liters"
    )

    set_cell_text(
        macro_table.rows[6].cells[1],
        f"{macros['weekly_calories']:,} kcal / week"
    )

    # Table 3 = 7-day meal matrix.
    meal_table = doc.tables[3]

    for day_index, pattern in enumerate(MEAL_PATTERNS):
        row = meal_table.rows[day_index + 1]

        for column_index, meal_index in enumerate(pattern, start=1):
            set_cell_text(
                row.cells[column_index],
                MEAL_LIBRARY[meal_index]
            )

        set_cell_text(
            row.cells[5],
            f"{macros['calories']:,} / {macros['protein']}g"
        )

    doc.save(output_docx)


# ============================================================
# 10. FILL WORKOUT DOCX
# ============================================================

def fill_workout_doc(data, output_docx):
    doc = Document(WORKOUT_TEMPLATE)

    replacements = {
        "[INSERT CLIENT NAME]": data["name"],
        "[INSERT GENDER]": data["gender"],
        "[INSERT AGE]": data["age"],
        "[INSERT HEIGHT]": data["height"],
        "[INSERT WEIGHT]": data["weight"],
        "[Fat Loss / Muscle Gain / Maintenance]": data["goal"],
        "[Beginner / Intermediate / Advanced]": data["experience"],
        "[Home / Gym / Both]": data["equipment"],
        "[Consistency]": data.get("obstacle", "Consistency"),
        "[INSERT INJURIES or 'None']": data.get("injuries", "None"),
        "[Strength / Fat Loss / General Fitness]":
            data.get("training_preference", "General Fitness"),
    }

    replace_everywhere(doc, replacements)

    # Table 1 = client snapshot.
    fill_client_snapshot(doc.tables[1], data)

    # Tables 4-10 = Day 1-Day 7.
    base_plan = GYM_PLAN if data["equipment"] == "Gym" else HOME_PLAN
    plan = adjust_volume(base_plan, data["experience"])

    for day_index in range(7):
        table = doc.tables[4 + day_index]
        exercises = plan[day_index]

        for exercise_index, exercise in enumerate(exercises, start=1):
            row = table.rows[exercise_index]

            for column_index, value in enumerate(exercise):
                set_cell_text(row.cells[column_index], value)

    doc.save(output_docx)


# ============================================================
# 11. DOCX -> PDF
# ============================================================

def find_libreoffice():
    for command in ("libreoffice", "soffice"):
        path = shutil.which(command)
        if path:
            return path

    common_paths = [
        "/usr/bin/libreoffice",
        "/usr/bin/soffice",
        "/usr/local/bin/libreoffice",
        "/usr/local/bin/soffice",
    ]

    for path in common_paths:
        if Path(path).exists():
            return path

    return None


def convert_to_pdf(docx_path, output_dir):
    libreoffice = find_libreoffice()

    if not libreoffice:
        raise RuntimeError(
            "LibreOffice/soffice is not installed on Render. "
            "Use the Dockerfile provided with this project."
        )

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    command = [
        libreoffice,
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(output_dir),
        str(docx_path),
    ]

    result = subprocess.run(
        command,
        capture_output=True,
        text=True,
        timeout=120,
    )

    if result.returncode != 0:
        raise RuntimeError(
            result.stderr.strip()
            or result.stdout.strip()
            or "LibreOffice PDF conversion failed."
        )

    pdf_path = output_dir / f"{Path(docx_path).stem}.pdf"

    if not pdf_path.exists():
        raise RuntimeError(
            f"LibreOffice reported success, but PDF was not found: {pdf_path}"
        )

    return pdf_path


# ============================================================
# 12. GENERATE CLIENT PACKAGE
# ============================================================

def safe_filename(name):
    cleaned = re.sub(r"[^A-Za-z0-9_\- ]+", "", name)
    cleaned = cleaned.strip().replace(" ", "_")

    return cleaned or "Client"


def generate_client_package(data):
    macros = calculate_macros(data)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    client_name = safe_filename(data["name"])

    client_folder = OUTPUT_DIR / f"{client_name}_{timestamp}"
    client_folder.mkdir(parents=True, exist_ok=True)

    workout_docx = (
        client_folder /
        f"{client_name}_Workout_Blueprint.docx"
    )

    nutrition_docx = (
        client_folder /
        f"{client_name}_Nutrition_Blueprint.docx"
    )

    fill_workout_doc(data, workout_docx)
    fill_nutrition_doc(data, macros, nutrition_docx)

    workout_pdf = convert_to_pdf(
        workout_docx,
        client_folder
    )

    nutrition_pdf = convert_to_pdf(
        nutrition_docx,
        client_folder
    )

    return macros, workout_pdf, nutrition_pdf


# ============================================================
# 13. TELEGRAM QUESTIONS
# ============================================================

QUESTIONS = [
    ("name", "👤 Client full name?"),

    (
        "gender",
        "♂️/♀️ Gender?\n\nReply: Male or Female"
    ),

    ("age", "🎂 Age?"),

    ("height", "📏 Height in cm?\n\nExample: 175"),

    ("weight", "⚖️ Weight in kg?\n\nExample: 70"),

    (
        "goal",
        "🎯 Primary goal?\n\n"
        "Reply: Fat Loss, Muscle Gain, or Maintenance"
    ),

    (
        "activity",
        "🏃 Activity level?\n\n"
        "Reply: Sedentary, Lightly Active, Moderate, or Very Active"
    ),

    (
        "experience",
        "🏋️ Training experience?\n\n"
        "Reply: Beginner, Intermediate, or Advanced"
    ),

    (
        "equipment",
        "🏠 Equipment available?\n\n"
        "Reply: Home or Gym"
    ),

    (
        "training_preference",
        "💪 Training preference?\n\n"
        "Reply: Strength, Fat Loss, or General Fitness"
    ),

    (
        "obstacle",
        "🧠 Main obstacle?\n\n"
        "Example: Consistency"
    ),

    (
        "restrictions",
        "🍽️ Diet restrictions/allergies?\n\n"
        "Reply None if there are none."
    ),

    (
        "injuries",
        "⚠️ Health conditions/injuries?\n\n"
        "Reply None if there are none."
    ),

    (
        "meals",
        "🍴 Meals per day?\n\n"
        "Usually 4. Reply with a number."
    ),
]


# ============================================================
# 14. TELEGRAM COMMANDS
# ============================================================

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update):
        return await deny(update)

    context.user_data.clear()
    context.user_data["step"] = 0

    await update.message.reply_text(
        "🧑‍💻 SIMON ORIGIN TRANSFORMATION\n"
        "PRIVATE PLAN GENERATOR\n\n"
        "You enter the client's information.\n"
        "I generate the personalized Workout + Nutrition PDFs.\n\n"
        "Let's begin.\n\n"
        + QUESTIONS[0][1]
    )


async def generate(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update):
        return await deny(update)

    context.user_data.clear()
    context.user_data["step"] = 0

    await update.message.reply_text(
        "🆕 NEW CLIENT PLAN\n\n"
        + QUESTIONS[0][1]
    )


async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update):
        return await deny(update)

    context.user_data.clear()

    await update.message.reply_text(
        "❌ Generation cancelled.\n\n"
        "Use /generate when you're ready."
    )


async def status(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update):
        return await deny(update)

    libreoffice = find_libreoffice()

    await update.message.reply_text(
        "🩺 SIMON ORIGIN SYSTEM STATUS\n\n"
        f"Excel calculator: "
        f"{'✅ OK' if CALCULATOR_FILE.exists() else '❌ MISSING'}\n"
        f"Workout template: "
        f"{'✅ OK' if WORKOUT_TEMPLATE.exists() else '❌ MISSING'}\n"
        f"Nutrition template: "
        f"{'✅ OK' if NUTRITION_TEMPLATE.exists() else '❌ MISSING'}\n"
        f"LibreOffice: "
        f"{'✅ OK' if libreoffice else '❌ MISSING'}\n"
        f"Admin ID: {ADMIN_ID}\n\n"
        "Use /generate to create a plan."
    )


# ============================================================
# 15. INPUT VALIDATION
# ============================================================

def normalize_choice(value, allowed):
    cleaned = value.strip().lower()

    for item in allowed:
        if cleaned == item.lower():
            return item

    return None


def normalize_number(value):
    value = value.strip()

    try:
        number = float(value)

        if number.is_integer():
            return str(int(number))

        return str(number)

    except ValueError:
        return None


# ============================================================
# 16. HANDLE ADMIN ANSWERS
# ============================================================

async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update):
        return await deny(update)

    if "step" not in context.user_data:
        await update.message.reply_text(
            "Use /generate to create a new client plan."
        )
        return

    step = context.user_data["step"]

    if step >= len(QUESTIONS):
        context.user_data.clear()
        await update.message.reply_text(
            "Use /generate to start a new plan."
        )
        return

    key, _ = QUESTIONS[step]
    value = update.message.text.strip()

    # -------------------------
    # NAME
    # -------------------------

    if key == "name":
        if len(value) < 2:
            await update.message.reply_text(
                "Please enter the client's full name."
            )
            return

    # -------------------------
    # GENDER
    # -------------------------

    elif key == "gender":
        value = normalize_choice(
            value,
            ["Male", "Female"]
        )

        if not value:
            await update.message.reply_text(
                "Please reply exactly: Male or Female"
            )
            return

    # -------------------------
    # AGE
    # -------------------------

    elif key == "age":
        try:
            age = int(value)

            if age < 18 or age > 100:
                raise ValueError

            value = str(age)

        except ValueError:
            await update.message.reply_text(
                "Please enter an age between 18 and 100."
            )
            return

    # -------------------------
    # HEIGHT
    # -------------------------

    elif key == "height":
        value = normalize_number(value)

        if value is None:
            await update.message.reply_text(
                "Enter height in cm. Example: 175"
            )
            return

        if not 100 <= float(value) <= 250:
            await update.message.reply_text(
                "Height should be between 100 and 250 cm."
            )
            return

    # -------------------------
    # WEIGHT
    # -------------------------

    elif key == "weight":
        value = normalize_number(value)

        if value is None:
            await update.message.reply_text(
                "Enter weight in kg. Example: 70"
            )
            return

        if not 30 <= float(value) <= 300:
            await update.message.reply_text(
                "Weight should be between 30 and 300 kg."
            )
            return

    # -------------------------
    # GOAL
    # -------------------------

    elif key == "goal":
        value = normalize_choice(
            value,
            list(GOAL_ADJUSTMENTS.keys())
        )

        if not value:
            await update.message.reply_text(
                "Reply: Fat Loss, Muscle Gain, or Maintenance"
            )
            return

    # -------------------------
    # ACTIVITY
    # -------------------------

    elif key == "activity":
        value = normalize_choice(
            value,
            list(ACTIVITY_MULTIPLIERS.keys())
        )

        if not value:
            await update.message.reply_text(
                "Reply: Sedentary, Lightly Active, Moderate, "
                "or Very Active"
            )
            return

    # -------------------------
    # EXPERIENCE
    # -------------------------

    elif key == "experience":
        value = normalize_choice(
            value,
            ["Beginner", "Intermediate", "Advanced"]
        )

        if not value:
            await update.message.reply_text(
                "Reply: Beginner, Intermediate, or Advanced"
            )
            return

    # -------------------------
    # EQUIPMENT
    # -------------------------

    elif key == "equipment":
        value = normalize_choice(
            value,
            ["Home", "Gym"]
        )

        if not value:
            await update.message.reply_text(
                "Reply: Home or Gym"
            )
            return

    # -------------------------
    # TRAINING PREFERENCE
    # -------------------------

    elif key == "training_preference":
        value = normalize_choice(
            value,
            ["Strength", "Fat Loss", "General Fitness"]
        )

        if not value:
            await update.message.reply_text(
                "Reply: Strength, Fat Loss, or General Fitness"
            )
            return

    # -------------------------
    # MEALS
    # -------------------------

    elif key == "meals":
        try:
            meal_count = int(value)

            if not 2 <= meal_count <= 6:
                raise ValueError

            value = f"{meal_count} meals/day"

        except ValueError:
            await update.message.reply_text(
                "Enter a number from 2 to 6."
            )
            return

    # -------------------------
    # SAVE ANSWER
    # -------------------------

    context.user_data[key] = value
    context.user_data["step"] += 1

    next_step = context.user_data["step"]

    if next_step < len(QUESTIONS):
        await update.message.reply_text(
            QUESTIONS[next_step][1]
        )
        return

    # ========================================================
    # ALL INFORMATION COLLECTED
    # ========================================================

    data = {
        key: context.user_data[key]
        for key, _ in QUESTIONS
    }

    await update.message.reply_text(
        "⏳ BUILDING PLAN...\n\n"
        "🧮 Calculating nutrition targets\n"
        "🏋️ Building workout\n"
        "🍽️ Building meal matrix\n"
        "📄 Filling templates\n"
        "🔄 Converting to PDF\n\n"
        "Please wait."
    )

    try:
        macros, workout_pdf, nutrition_pdf = (
            generate_client_package(data)
        )

        # Send summary first.
        await update.message.reply_text(
            "✅ PLAN GENERATED\n\n"
            f"Client: {data['name']}\n\n"
            f"🔥 Calories: {macros['calories']:,} kcal/day\n"
            f"🥩 Protein: {macros['protein']} g\n"
            f"🍚 Carbs: {macros['carbs']} g\n"
            f"🥑 Fat: {macros['fats']} g\n"
            f"💧 Water: {macros['water']} L\n\n"
            "⚠️ Review both PDFs before sending them to the client."
        )

        # Workout PDF.
        with open(workout_pdf, "rb") as file:
            await update.message.reply_document(
                document=file,
                caption=(
                    f"🏋️ {data['name']}\n"
                    "30-Day Workout Blueprint"
                ),
            )

        # Nutrition PDF.
        with open(nutrition_pdf, "rb") as file:
            await update.message.reply_document(
                document=file,
                caption=(
                    f"🍽️ {data['name']}\n"
                    "30-Day Nutrition Blueprint"
                ),
            )

    except Exception as error:
        await update.message.reply_text(
            "❌ PLAN GENERATION FAILED\n\n"
            f"{error}\n\n"
            "Use /status to check the required files and "
            "LibreOffice."
        )

    finally:
        context.user_data.clear()


# ============================================================
# 17. MAIN
# ============================================================

def main():
    application = (
        Application
        .builder()
        .token(BOT_TOKEN)
        .build()
    )

    application.add_handler(
        CommandHandler("start", start)
    )

    application.add_handler(
        CommandHandler("generate", generate)
    )

    application.add_handler(
        CommandHandler("cancel", cancel)
    )

    application.add_handler(
        CommandHandler("status", status)
    )

    application.add_handler(
        MessageHandler(
            filters.TEXT & ~filters.COMMAND,
            handle_text,
        )
    )

    print("Simon Origin Plan Generator is running...")

    application.run_polling(
        allowed_updates=Update.ALL_TYPES
    )


if __name__ == "__main__":
    main()
